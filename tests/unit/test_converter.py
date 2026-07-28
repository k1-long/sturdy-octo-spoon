"""converter.py 单元测试 — 测试单个组件的转换功能"""

from io import BytesIO

import pytest
from docx import Document as DocxDocument
from src.converter import markdown_to_html, markdown_to_docx, validate_markdown


def _read_docx(buffer):
    """读取 docx BytesIO 返回 Document 对象"""
    buffer.seek(0)
    return DocxDocument(buffer)


def _get_texts(doc):
    return [p.text for p in doc.paragraphs]


class TestMarkdownToHtml:
    """1️⃣ 单元测试：HTML 转换"""

    def test_basic_conversion(self):
        """测试基本 Markdown 转 HTML"""
        result = markdown_to_html("# Hello\n\nWorld")
        assert "<h1>Hello</h1>" in result
        assert "<p>World</p>" in result

    def test_code_block(self):
        """测试代码块转换"""
        md = "```python\ndef hello():\n    pass\n```"
        result = markdown_to_html(md)
        assert "<code>" in result

    def test_table(self):
        """测试表格转换"""
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        result = markdown_to_html(md)
        assert "<table>" in result


class TestMarkdownToDocx:
    """1️⃣ 单元测试：Word 文档转换 — 验证内容而非仅大小"""

    def test_basic_conversion(self):
        """标题和段落正确转换"""
        buffer = markdown_to_docx("# Test\n\nHello World")
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        # 标题在 Word 中以 heading 段落存在
        assert any("Test" in t for t in texts)
        assert any("Hello World" in t for t in texts)
        buffer.close()

    def test_all_heading_levels(self):
        """所有标题层级 (h1-h6)"""
        md = "# H1\n## H2\n### H3"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        assert any("H1" in t for t in texts)
        assert any("H2" in t for t in texts)
        assert any("H3" in t for t in texts)
        buffer.close()

    def test_unicode_content(self):
        """3️⃣ 边界条件：Unicode 和表情符号"""
        buffer = markdown_to_docx("# Unicode 测试 🎉\n\n中文内容")
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        assert any("Unicode" in t for t in texts)
        assert any("中文" in t for t in texts)
        buffer.close()

    def test_large_content(self):
        """3️⃣ 边界条件：超长文本"""
        long_text = "# Test\n\n" + "Hello " * 1000
        buffer = markdown_to_docx(long_text)
        doc = _read_docx(buffer)
        assert len(doc.paragraphs) >= 2
        buffer.close()

    def test_unordered_list(self):
        """无序列表"""
        md = "- item 1\n- item 2\n- item 3"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        assert any("item 1" in t for t in texts)
        assert any("item 2" in t for t in texts)
        assert any("item 3" in t for t in texts)
        buffer.close()

    def test_ordered_list(self):
        """有序列表"""
        md = "1. first\n2. second\n3. third"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        assert any("first" in t for t in texts)
        assert any("second" in t for t in texts)
        buffer.close()

    def test_table(self):
        """表格"""
        md = "| A | B |\n|---|---|\n| 1 | 2 |"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        tables = doc.tables
        assert len(tables) >= 1
        # 验证表格内容
        cell_text = tables[0].cell(0, 0).text
        assert "A" in cell_text or "1" in cell_text
        buffer.close()

    def test_blockquote(self):
        """引用块"""
        md = "> quoted text"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        assert any("quoted text" in t for t in texts)
        buffer.close()

    def test_bold_and_italic_inline(self):
        """粗体和斜体内联元素"""
        md = "**bold** and *italic* and `code`"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        full_text = " ".join(texts)
        assert "bold" in full_text
        assert "italic" in full_text
        assert "code" in full_text
        buffer.close()

    def test_nested_bold_italic(self):
        """嵌套内联格式：粗体中的斜体"""
        md = "**bold *and italic* end**"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        full_text = " ".join(texts)
        assert "bold" in full_text
        assert "italic" in full_text
        buffer.close()

    def test_pre_without_code_tag(self):
        """pre 标签内无 code 子标签"""
        buffer = markdown_to_docx("    indented code block")
        doc = _read_docx(buffer)
        assert len(doc.paragraphs) > 0
        buffer.close()

    def test_empty_content_generates_valid_docx(self):
        """空 Markdown 仍生成有效 docx 容器"""
        buffer = markdown_to_docx("")
        doc = _read_docx(buffer)
        assert doc is not None
        buffer.close()

    def test_bold_italic_in_heading(self):
        """标题中的粗体和斜体"""
        md = "# **Bold** and *Italic* Heading"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        full_text = " ".join(texts)
        assert "Bold" in full_text
        assert "Italic" in full_text
        buffer.close()

    def test_horizontal_rule(self):
        """分割线"""
        md = "before\n\n---\n\nafter"
        buffer = markdown_to_docx(md)
        doc = _read_docx(buffer)
        texts = _get_texts(doc)
        full_text = " ".join(texts)
        assert "before" in full_text
        assert "after" in full_text
        buffer.close()


class TestValidateMarkdown:
    """验证函数测试"""

    def test_valid_input(self):
        """正常输入无问题"""
        issues = validate_markdown("# Valid markdown")
        assert issues == []

    def test_empty_input(self):
        """3️⃣ 边界条件：空输入"""
        issues = validate_markdown("")
        assert len(issues) == 1
        assert "为空" in issues[0]

    def test_whitespace_only(self):
        """3️⃣ 边界条件：仅空白字符"""
        issues = validate_markdown("   \n  \t  ")
        assert len(issues) == 1

    def test_non_string_input(self):
        """4️⃣ 错误恢复：非字符串类型输入"""
        issues = validate_markdown(123)
        assert len(issues) == 1
        assert "字符串" in issues[0]

    def test_none_input(self):
        """4️⃣ 错误恢复：None 输入"""
        issues = validate_markdown(None)
        assert len(issues) >= 1
