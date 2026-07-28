"""Markdown 转 Word 文档的核心转换模块"""

from io import BytesIO

from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import markdown


def markdown_to_html(md_text: str) -> str:
    """将 Markdown 文本转换为 HTML"""
    return markdown.markdown(md_text, extensions=["extra", "codehilite", "tables"])


def markdown_to_docx(md_text: str) -> BytesIO:
    """将 Markdown 文本转换为 Word 文档（返回文件流）

    调用方负责关闭返回的 BytesIO。
    """
    html = markdown_to_html(md_text)
    soup = BeautifulSoup(html, "html.parser")
    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(11)

    _process_element(soup, doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def _process_element(node, doc, parent=None):
    """递归处理 HTML 节点树，写入 Word 文档"""
    if node.name is None:
        # 文本节点
        text = str(node).strip()
        if text and parent is not None:
            run = parent.add_run(text)
        elif text and parent is None:
            p = doc.add_paragraph()
            p.add_run(text)
        return

    tag = node.name

    if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
        level = int(tag[1])
        heading = doc.add_heading(level=level)
        for child in node.children:
            _process_child_inline(child, heading)

    elif tag == "p":
        p = doc.add_paragraph()
        for child in node.children:
            _process_child_inline(child, p)

    elif tag in ("ul", "ol"):
        for li in node.find_all("li", recursive=False):
            style_name = "List Bullet" if tag == "ul" else "List Number"
            p = doc.add_paragraph(style=style_name)
            for child in li.children:
                _process_child_inline(child, p)

    elif tag == "pre":
        for code_tag in node.find_all("code", recursive=False):
            code_text = code_tag.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        if not node.find("code"):
            code_text = node.get_text()
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)

    elif tag == "table":
        rows = node.find_all("tr")
        if rows:
            cols = len(rows[0].find_all(["th", "td"])) if rows else 1
            table = doc.add_table(rows=len(rows), cols=cols)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                cells = row.find_all(["th", "td"])
                for j, cell in enumerate(cells):
                    if j < cols:
                        # 保留内联格式：清空默认段落，逐子节点处理
                        cell_word = table.cell(i, j)
                        cell_word.text = ""
                        p = cell_word.paragraphs[0]
                        for child in cell.children:
                            _process_child_inline(child, p)

    elif tag in ("blockquote",):
        # 引用块：添加缩进和斜体以在视觉上区分
        for child in node.children:
            if child.name:
                _add_blockquote_paragraph(child, doc)

    elif tag == "hr":
        # 分割线：添加带底部边框的空段落
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.space_before = Pt(6)
        pf.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = pPr.makeelement(qn("w:pBdr"), {})
        bottom = pBdr.makeelement(qn("w:bottom"), {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "999999",
        })
        pBdr.append(bottom)
        pPr.append(pBdr)

    elif tag == "img":
        # 图片：添加占位文字（python-docx 插入图片需要文件流，此处用 alt 文本替代）
        alt = node.get("alt", "")
        src = node.get("src", "")
        p = doc.add_paragraph()
        run = p.add_run(f"[图片: {alt}]" if alt else "[图片]")
        run.italic = True
        run.font.color.rgb = None  # 灰色

    else:
        for child in node.children:
            _process_element(child, doc, parent)


def _process_child_inline(node, paragraph, **styles):
    """递归处理段落/标题内的内联元素（粗体、斜体、代码、链接、纯文本），支持嵌套格式

    通过 styles 参数传递累积的样式（bold/italic/underline/font_name/font_size）。
    """
    if node.name is None:
        text = str(node)
        if text.strip():
            run = paragraph.add_run(text)
            if styles.get("bold"):
                run.bold = True
            if styles.get("italic"):
                run.italic = True
            if styles.get("underline"):
                run.underline = True
            fn = styles.get("font_name")
            if fn:
                run.font.name = fn
            fs = styles.get("font_size")
            if fs:
                run.font.size = fs
        return

    tag = node.name

    # 合并父级传递的样式与当前标签的样式
    kw = dict(styles)
    if tag in ("strong", "b"):
        kw["bold"] = True
    elif tag in ("em", "i"):
        kw["italic"] = True
    elif tag == "code":
        kw["font_name"] = "Consolas"
        kw["font_size"] = Pt(9.5)
    elif tag == "a":
        kw["underline"] = True
    else:
        for child in node.children:
            _process_child_inline(child, paragraph, **kw)
        return

    for child in node.children:
        _process_child_inline(child, paragraph, **kw)


def _add_blockquote_paragraph(node, doc):
    """将 blockquote 内的元素以斜体 + 左缩进写入文档，使其在视觉上可区分"""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.left_indent = Pt(24)
    for child in node.children:
        _process_child_inline(child, p, italic=True)


def validate_markdown(md_text: str) -> list[str]:
    """验证 Markdown 文本，返回问题列表（空列表表示无问题）"""
    issues = []
    if not isinstance(md_text, str):
        issues.append("输入必须是字符串类型")
        return issues
    if not md_text.strip():
        issues.append("输入内容为空")
    return issues
